# -*- coding: utf-8 -*-
"""test_file 全量模板填报复测（字段提取 + 映射 + Word 填充校验）。

用法:
  py scripts/verify_template_fill_batch.py
  py scripts/verify_template_fill_batch.py test_sample/test_file/2014-兴泰贸易.pdf
  py scripts/verify_template_fill_batch.py --no-word   # 仅检查字段映射，不跑 Word COM
"""
from __future__ import annotations

import json
import os
import re
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from field_mapping import expand_fields_for_template

TEST_DIR = "test_sample/test_file"

# 模板占位符 → 不应出现的脏值特征
SLOT_RULES = {
    "档案卷宗": {
        "委托代理合同中委托人": lambda v: v and "原告" not in v and "被告" not in v and "待确认" not in v,
        "判决书中的原告": lambda v: v and "原告" not in v[:4] and "待确认" not in v,
        "判决书中的被告": lambda v: v and "被告" not in v[:4] and "待确认" not in v,
        "判决书原告的委托诉讼代理人": lambda v: v and v not in ("无", "待确认", "律师"),
    },
    "送达材料清单": {
        "判决书上的原告": lambda v: v and "原告" not in v[:4] and "，" not in v,
        "判决书上代理律师": lambda v: v and v not in ("无", "待确认"),
    },
    "结案报告表": {
        "结案小结": lambda v: v and len(v) >= 10 and "待确认" not in v,
        "审（办）结果": lambda v: v and len(v) >= 10,
    },
    "立案审批表": {
        "3": lambda v: v and "原告" not in v and "待确认" not in v,
        "4": lambda v: v and "原告" not in v[:4] and "被告" not in v,
    },
}

FIELD_CRITICAL = [
    "委托人", "判决书中的原告", "判决书中的被告",
    "承办律师", "判决书上代理律师", "案由", "审理法院",
]


def _silent(*a, **k):
    pass


def _check_mapped(template_name: str, fields: dict) -> list[str]:
    mapped = expand_fields_for_template(template_name, fields)
    bad = []
    rules = SLOT_RULES.get(template_name, {})
    for ph, pred in rules.items():
        val = (mapped.get(ph) or "").strip()
        if not val:
            bad.append(f"{template_name}.{ph}=空")
        elif not pred(val):
            bad.append(f"{template_name}.{ph}={val[:40]!r}")
    return bad


def _scan_docx_placeholders(docx_path: str) -> list[str]:
    try:
        from docx import Document
    except ImportError:
        return []
    issues = []
    doc = Document(docx_path)
    text = ""
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                text += c.text
    for p in doc.paragraphs:
        text += p.text
    left = re.findall(r"【([^】]{1,50})】", text)
    key_left = [x for x in left if any(k in x for k in ("委托人", "原告", "被告", "律师", "案由", "法院"))]
    if key_left:
        issues.append(f"残留占位符: {', '.join(sorted(set(key_left))[:6])}")
    if re.search(r"【[^】]*待确认[^】]*】", text):
        issues.append("残留【待确认】占位符")
    elif re.search(
        r"(委托人|当事人|对方当事人|原告|被告|承办律师|委托方)[：:\s]*待确认",
        text,
    ):
        issues.append("关键标签仍为待确认")
    if re.search(r"原告[：:].*被告", text):
        issues.append("整串当事人未拆分")
    return issues


def analyze_one(pdf_path: str, config: dict, *, run_word: bool) -> dict:
    from archive_pipeline import analyze_archive

    bn = os.path.basename(pdf_path)
    log_fn = _silent if not os.environ.get("VERBOSE") else print

    if run_word:
        analysis = analyze_archive("civil", pdf_path, config, log=log_fn)
        fields = analysis.fields or {}
        tpl_issues = list(analysis.template_issues or [])
        docx_paths = dict(analysis.generated_templates or {})
    else:
        from archive_pipeline import (
            ingest_archive_sources,
            segment_and_map_documents,
            extract_fields_auto,
            normalize_fields,
        )
        from document_segmenter import DocumentSource, DOC_TYPE_DEFAULT, build_segmented_from_units

        doc_sources = [DocumentSource(path=pdf_path, doc_type=DOC_TYPE_DEFAULT)]
        pdf_texts, page_texts, layout_blocks, _, _ = ingest_archive_sources(
            doc_sources, config, log=_silent
        )
        doc_spans = segment_and_map_documents(
            doc_sources, "civil", config,
            pdf_texts_by_path=pdf_texts,
            page_texts_by_path=page_texts,
            layout_blocks_by_path=layout_blocks,
            log=_silent,
        )
        pdf_text = "\n".join(pdf_texts.values())
        segmented = build_segmented_from_units(doc_spans, page_texts)
        raw = extract_fields_auto(pdf_text, segmented=segmented, log=log_fn)
        fields = normalize_fields(raw, pdf_text)
        tpl_issues = []
        docx_paths = {}

    slot_issues = []
    for tpl in ("档案卷宗", "送达材料清单", "立案审批表"):
        slot_issues.extend(_check_mapped(tpl, fields))

    missing_fields = [
        f for f in FIELD_CRITICAL
        if not (fields.get(f) or "").strip()
        or str(fields.get(f, "")).strip() in ("待确认", "无")
    ]

    docx_issues = []
    for name, path in docx_paths.items():
        if os.path.isfile(path):
            docx_issues.extend(f"{name}: {x}" for x in _scan_docx_placeholders(path))

    ok = not slot_issues and not missing_fields and not tpl_issues and not docx_issues
    return {
        "case": bn,
        "ok": ok,
        "fields_n": len(fields),
        "委托人": (fields.get("委托人") or "")[:30],
        "原告": (fields.get("判决书中的原告") or fields.get("原告") or "")[:30],
        "被告": (fields.get("判决书中的被告") or "")[:30],
        "律师": (fields.get("判决书上代理律师") or fields.get("承办律师") or "")[:20],
        "missing_fields": missing_fields,
        "slot_issues": slot_issues,
        "template_issues": tpl_issues[:5],
        "docx_issues": docx_issues[:5],
    }


def main():
    import argparse
    from settings import load_config

    ap = argparse.ArgumentParser()
    ap.add_argument("pdfs", nargs="*", help="指定 PDF，默认 test_file 全量")
    ap.add_argument("--no-word", action="store_true", help="跳过 Word 填充，仅字段+映射")
    args = ap.parse_args()

    config = load_config()
    pdfs = args.pdfs or sorted(
        os.path.join(TEST_DIR, f)
        for f in os.listdir(TEST_DIR)
        if f.endswith(".pdf")
    )
    pdfs = [p for p in pdfs if os.path.isfile(p)]
    if not pdfs:
        print("无测试 PDF")
        sys.exit(1)

    results = []
    failed = 0
    for i, p in enumerate(pdfs, 1):
        print(f"[{i}/{len(pdfs)}] {os.path.basename(p)} ...", flush=True)
        try:
            r = analyze_one(p, config, run_word=not args.no_word)
        except Exception as e:
            r = {"case": os.path.basename(p), "ok": False, "error": str(e)}
        results.append(r)
        if not r.get("ok"):
            failed += 1
            parts = []
            if r.get("error"):
                parts.append(f"ERR={r['error']}")
            if r.get("missing_fields"):
                parts.append(f"缺字段={r['missing_fields']}")
            if r.get("slot_issues"):
                parts.append(f"槽位={r['slot_issues'][:3]}")
            if r.get("template_issues"):
                parts.append(f"模板={r['template_issues']}")
            if r.get("docx_issues"):
                parts.append(f"docx={r['docx_issues']}")
            print(f"  FAIL {' | '.join(parts)}", flush=True)
        else:
            print(
                f"  OK 委托人={r.get('委托人')} 原告={r.get('原告')} 律师={r.get('律师')}",
                flush=True,
            )

    out = "outputs/_template_fill_batch.json"
    os.makedirs("outputs", exist_ok=True)
    with open(out, "w", encoding="utf-8") as f:
        json.dump(results, f, ensure_ascii=False, indent=2)

    ok_n = sum(1 for r in results if r.get("ok"))
    print(f"\n==== {ok_n}/{len(results)} 通过, {failed} 失败 ====")
    print(f"saved {out}")
    sys.exit(1 if failed else 0)


if __name__ == "__main__":
    main()
