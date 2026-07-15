#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
归档流水线 V2
输入：一份案件档案 PDF
输出：填好的 5 份 docx（默认，见 output.docx_only）

V4 新增：analyze_archive / assemble_archive 完整归档流程
"""

import os
import re
import json
import zipfile
import shutil
import threading
from dataclasses import dataclass, field
from datetime import datetime
from typing import Dict, List, Optional, Set

from field_mapping import get_template_paths, expand_fields_for_template
from template_filler import TemplateFiller
from settings import (
    load_config,
    get_deepseek_config,
    get_ocr_engine,
    get_extraction_mode,
    get_archive_order_mode,
    output_docx_only,
    parse_llm_output,
)
from archive_ocr import extract_pdf_text, get_pdf_page_count
from case_outcome import (
    build_outcome_from_units,
    detect_outcome_warnings,
    ensure_outcome_covers_execution,
    unify_case_outcome_fields,
)
from pdf_text_chunk import build_pdf_chunk_for_llm
from document_segmenter import (
    DOC_TYPE_COMPLAINT,
    DOC_TYPE_CONTRACT,
    DOC_TYPE_DEFAULT,
    DOC_TYPE_EXECUTION,
    DOC_TYPE_INDICTMENT,
    DOC_TYPE_JUDGMENT,
    DOC_TYPE_POA,
    DOC_TYPE_LABELS,
    DOC_TYPE_OTHER,
    DocumentSource,
    build_segmented_from_units,
    build_segmented_text,
    validate_sources_for_archive,
)
from field_merger import merge_partial_fields
from field_sanitize import (
    sanitize_all_field_values,
    sanitize_court_case_no,
    parse_court_document_list,
)

# V4 新增导入
try:
    import archive_catalog as ac
    import pdf_doc_locator as pdl
    import page_ocr
except ImportError:
    pass  # V4 模块可能未安装

try:
    from app_paths import get_outputs_dir, get_prompt_path
except ImportError:

    def get_outputs_dir():
        d = os.path.join(os.path.dirname(__file__), "outputs")
        os.makedirs(d, exist_ok=True)
        return d

    def get_prompt_path():
        return os.path.join(os.path.dirname(__file__), "prompts", "extract_prompt.txt")


def _prompts_dir():
    try:
        from app_paths import get_app_dir

        return os.path.join(get_app_dir(), "prompts")
    except ImportError:
        return os.path.join(os.path.dirname(__file__), "prompts")


def load_prompt_file(filename: str) -> str:
    path = os.path.join(_prompts_dir(), filename)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        if "文档内容" not in text:
            text = text.rstrip() + "\n\n文档内容：\n"
        return text
    raise FileNotFoundError(f"提示词文件不存在: {path}")


def load_extract_prompt(prompt_path=None):
    path = prompt_path or get_prompt_path()
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
        if "文档内容" not in text:
            text = text.rstrip() + "\n\n文档内容：\n"
        return text
    raise FileNotFoundError(f"提示词文件不存在: {path}")


CASE_TYPE_LABELS = {
    "civil": "民事",
    "criminal": "刑事",
    "admin": "行政",
    "nonlit": "非诉",
    "counsel": "法律顾问",
}


def detect_case_type(pdf_text: str) -> str:
    """Best-effort case type detection used only when callers omit case_type."""
    text = (pdf_text or "")[:30000]
    def weighted(strong: tuple[str, ...], weak: tuple[str, ...] = ()) -> int:
        return 6 * sum(text.count(x) for x in strong) + sum(text.count(x) for x in weak)

    scores = {
        "criminal": weighted(
            ("刑事判决书", "刑事裁定书", "刑事起诉书", "公诉书"),
            ("公诉机关", "被告人", "犯罪嫌疑人", "辩护人", "罪名"),
        ),
        "admin": weighted(
            ("行政判决书", "行政裁定书", "行政起诉状", "行政复议决定书"),
            ("行政诉讼", "被诉行政行为", "行政机关", "行政处罚"),
        ),
        "counsel": weighted(
            ("常年法律顾问合同", "法律顾问合同", "常年法律顾问"),
            ("顾问单位", "法律咨询", "合同审查"),
        ),
        "nonlit": weighted(
            ("非诉讼法律事务委托合同", "专项法律服务合同", "法律尽职调查报告"),
            ("非诉讼法律事务", "专项法律服务", "尽职调查"),
        ),
        "civil": weighted(
            ("民事判决书", "民事裁定书", "民事起诉状", "民事调解书"),
            ("原告", "被告", "上诉人", "被上诉人"),
        ),
    }
    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else "civil"


def resolve_case_type(case_type: Optional[str], pdf_text: str = "") -> str:
    return case_type if case_type in CASE_TYPE_LABELS else detect_case_type(pdf_text)


def load_case_prompt(case_type: str) -> str:
    filename = f"extract_{case_type}.txt"
    path = os.path.join(_prompts_dir(), filename)
    return load_prompt_file(filename) if os.path.exists(path) else load_extract_prompt()


_CRITICAL_FIELDS = ["委托人", "当事人", "对方当事人", "案由", "审理法院", "承办律师"]


def _deepseek_request(user_content: str, system_content: str) -> str:
    """单次 DeepSeek 调用，返回原始回复文本。"""
    import requests

    ds = get_deepseek_config()
    if not ds.get("api_key"):
        raise RuntimeError("请在 config.json 中配置 deepseek.api_key")
    headers = {
        "Authorization": f"Bearer {ds['api_key']}",
        "Content-Type": "application/json",
    }
    data = {
        "model": ds["model"],
        "messages": [
            {"role": "system", "content": system_content},
            {"role": "user", "content": user_content},
        ],
        "temperature": 0,
    }
    response = requests.post(
        ds["base_url"] + "/chat/completions",
        headers=headers,
        json=data,
        timeout=180,
    )
    result = response.json()
    if "choices" not in result:
        raise RuntimeError(f"DeepSeek API 错误: {result}")
    return result["choices"][0]["message"]["content"]


def _deepseek_chat(user_content: str, system_content: str, *, max_retries: int = 1) -> dict:
    """调用 DeepSeek 并解析为字段字典。

    解析为空（总体失败）时最多重试 ``max_retries`` 次，附带更明确的输出格式要求，
    降低偶发的非结构化/空输出概率。
    """
    last_err = None
    attempt = 0
    extra_hint = ""
    while attempt <= max_retries:
        attempt += 1
        try:
            content = _deepseek_request(user_content + extra_hint, system_content)
            parsed = parse_llm_output(content)
        except Exception as e:  # 网络/接口异常：重试
            last_err = e
            if attempt > max_retries:
                raise RuntimeError(f"LLM 调用失败: {e}")
            extra_hint = ""
            continue

        if parsed:
            from field_sanitize import is_valid_field_value

            cleaned = {}
            for k, v in parsed.items():
                sv = "" if v is None else str(v).strip()
                if is_valid_field_value(sv):
                    cleaned[k] = sv
            parsed = cleaned
            empty_critical = [
                f for f in _CRITICAL_FIELDS
                if not parsed.get(f) or str(parsed[f]).strip() in ("", "待确认")
            ]
            if empty_critical:
                print(f"  [WARN] LLM 返回的关键字段为空: {', '.join(empty_critical)}")
            return parsed

        # 解析为空：再给一次机会，强调输出格式
        if attempt <= max_retries:
            print("  [WARN] LLM 输出解析为空，重试一次（强调格式）")
            extra_hint = (
                "\n\n注意：请严格按「字段名: 值」逐行输出，"
                "找不到确切信息的字段直接留空（值为空），不要填「待确认」或说明文字。"
            )

    if last_err:
        raise RuntimeError(f"LLM 调用失败: {last_err}")
    return {}


def extract_fields_with_deepseek(pdf_text, prompt_path=None, case_type=None):
    case_type = resolve_case_type(case_type, pdf_text)
    case_label = CASE_TYPE_LABELS[case_type]
    system_content = (
        f"你是专业的法律文档分析助手，当前案件类型为{case_label}。"
        "只能依据文档原文提取，不得把其他案件类型的角色、案由或处理结果套入本案。"
        "结案小结应概括本案实际裁判、处理或服务结果，不超过150字；"
        "卷内存在执行材料时，必须综合全卷证据写明诉讼结果与执行措施、履行情况及最终执行结论。"
    )
    prompt_template = load_extract_prompt(prompt_path) if prompt_path else load_case_prompt(case_type)
    cfg = load_config()
    chunk = build_pdf_chunk_for_llm(pdf_text, ocr_engine=get_ocr_engine(cfg))
    return _deepseek_chat(prompt_template + chunk, system_content)


def extract_fields_segmented(segmented, log=print, case_type=None) -> dict:
    """分路提取：判决书 / 执行裁定书 / 委托代理合同"""
    combined = "\n".join(str(v or "") for v in segmented.values())
    case_type = resolve_case_type(case_type, combined)
    case_label = CASE_TYPE_LABELS[case_type]
    system_content = (
        f"你是专业的法律文档分析助手，当前案件类型为{case_label}。"
        "严格按提示词格式输出字段，不要额外说明；不得混用不同案件类型的当事人角色。"
        "执行分路必须通读全部执行阶段材料，以最终裁定为准并结合查控、查封、冻结、限制消费等事实综合归纳。"
    )
    partials = {}

    judgment_prompt = f"extract_judgment_{case_type}.txt"
    if not os.path.exists(os.path.join(_prompts_dir(), judgment_prompt)):
        judgment_prompt = "extract_judgment.txt"
    judgment_label = "刑事裁判文书" if case_type == "criminal" else "裁判文书"

    routes = [
        (DOC_TYPE_POA, "extract_poa.txt", "授权委托书"),
        (DOC_TYPE_JUDGMENT, judgment_prompt, judgment_label),
        (DOC_TYPE_CONTRACT, f"extract_contract_{case_type}.txt", "委托/法律服务合同"),
    ]
    if case_type == "criminal":
        routes.append((DOC_TYPE_INDICTMENT, "extract_criminal_indictment.txt", "起诉书/抗诉书"))
    if case_type in ("civil", "admin"):
        routes.append((DOC_TYPE_EXECUTION, "extract_execution.txt", "执行阶段材料"))
    processed_types = set()
    for doc_type, prompt_file, route_label in routes:
        text = segmented.get(doc_type)
        if not text or len(text.strip()) < 30:
            continue
        processed_types.add(doc_type)
        log(f"       分路提取：{route_label}（{len(text)} 字）")
        try:
            if prompt_file.startswith("extract_contract_") and not os.path.exists(os.path.join(_prompts_dir(), prompt_file)):
                prompt_file = "extract_contract.txt"
            prompt = load_prompt_file(prompt_file)
            chunk = text[:12000] if len(text) > 12000 else text
            partials[doc_type] = _deepseek_chat(prompt + chunk, system_content)
        except Exception as e:
            log(f"  [WARN] {route_label} 提取失败: {e}")

    complaint = (
        segmented.get(DOC_TYPE_COMPLAINT)
        or segmented.get(DOC_TYPE_OTHER)
        or segmented.get(DOC_TYPE_DEFAULT)
    )
    if complaint and len(complaint.strip()) >= 30:
        processed_types.update({DOC_TYPE_COMPLAINT, DOC_TYPE_OTHER, DOC_TYPE_DEFAULT})
        log(f"       分路提取：起诉状/其他（{len(complaint)} 字）")
        try:
            prompt = load_case_prompt(case_type)
            chunk = complaint[:8000]
            partials[DOC_TYPE_COMPLAINT] = _deepseek_chat(prompt + chunk, system_content)
        except Exception as e:
            log(f"  [WARN] 起诉状提取失败: {e}")

    # Non-litigation and counsel archives often consist of work records,
    # opinions, reports, or other document types that have no civil route.
    # Feed all remaining material to the case-specific prompt once.
    remaining = "\n\n".join(
        str(text or "")
        for doc_type, text in segmented.items()
        if doc_type not in processed_types and text and len(str(text).strip()) >= 30
    )
    if remaining:
        log(f"       分路提取：其他{case_label}材料（{len(remaining)} 字）")
        try:
            partials[DOC_TYPE_OTHER] = _deepseek_chat(
                load_case_prompt(case_type) + remaining[:12000],
                system_content,
            )
        except Exception as e:
            log(f"  [WARN] 其他材料提取失败: {e}")

    if not partials:
        return {}

    merged = merge_partial_fields(partials, case_type=case_type)
    log(f"       分路合并 {len(merged)} 个字段")
    return merged


def extract_fields_auto(pdf_text, segmented=None, log=print, case_type=None):
    """根据 config 选择 segmented 或 legacy 提取"""
    case_type = resolve_case_type(case_type, pdf_text)
    mode = get_extraction_mode()
    if mode == "segmented":
        if segmented is None:
            segmented_obj = build_segmented_text(pdf_text=pdf_text)
            segmented = segmented_obj.segments
        if segmented and any(segmented.values()):
            fields = extract_fields_segmented(segmented, log=log, case_type=case_type)
            if fields:
                return fields
            log("  [WARN] 分路提取无结果，回退 legacy")
    return extract_fields_with_deepseek(pdf_text, case_type=case_type)


def normalize_fields(raw, pdf_text="", doc_spans=None, page_texts_by_path=None, case_type=None):
    if not raw:
        return {}
    from field_sanitize import enrich_party_fields

    case_type = resolve_case_type(case_type, pdf_text)
    m = enrich_party_fields(dict(raw))
    m["案件类别"] = CASE_TYPE_LABELS[case_type]
    aliases = {
        "委托人名称": "委托人",
        "委托方": "当事人",
        "立案日期": "收案日期",
        "代理律师": "承办律师",
        # 新增：处理更多可能的字段名变体
        "原告": "当事人",
        "被告": "对方当事人",
        "被申请人": "对方当事人",
        "申请人": "当事人",
    }
    if case_type == "criminal":
        aliases.update({
            "被告人": "当事人",
            "犯罪嫌疑人": "当事人",
            "辩护人": "承办律师",
            "罪名": "案由",
            "公诉机关": "对方当事人",
            "审判法院": "审理法院",
        })
    elif case_type == "admin":
        aliases.update({
            "行政相对人": "当事人",
            "行政机关": "对方当事人",
            "被诉行政机关": "对方当事人",
            "审判法院": "审理法院",
            "被诉行政行为": "案情简介",
        })
    elif case_type == "nonlit":
        aliases.update({
            "项目委托人": "委托人",
            "项目事项": "案由",
            "项目类型": "案由",
            "项目律师": "承办律师",
            "服务成果": "结案小结",
        })
    elif case_type == "counsel":
        aliases.update({
            "顾问单位": "委托人",
            "顾问事项": "案由",
            "顾问律师": "承办律师",
            "服务成果": "结案小结",
        })
    for src, dst in aliases.items():
        if src in m and m[src] and (dst not in m or not m.get(dst)):
            m[dst] = m[src]
    if m.get("委托人联系地址及电话") and not m.get("委托人电话"):
        m["委托人电话"] = m["委托人联系地址及电话"]
    if case_type in ("nonlit", "counsel") and m.get("结案小结") and not m.get("审（办）结果"):
        m["审（办）结果"] = m["结案小结"]
    from field_mapping import _build_case_project_name

    if case_type == "criminal":
        accused = str(m.get("被告人") or m.get("犯罪嫌疑人") or m.get("当事人") or "").strip()
        charge = str(m.get("罪名") or m.get("案由") or "").strip()
        full_case_name = str(m.get("案件或项目名称") or "").strip()
        if not full_case_name and (accused or charge):
            full_case_name = f"被告人{accused}{'涉嫌' if charge and not charge.endswith('罪') else ''}{charge}一案"
    else:
        full_case_name = _build_case_project_name(m)
    if full_case_name:
        m["案件或项目名称"] = full_case_name
    if m.get("案由") and not m.get("案件或项目名称"):
        m["案件或项目名称"] = _build_case_project_name(m)
    elif m.get("案件或项目名称") and not m.get("案由"):
        ay = m["案件或项目名称"]
        if "诉" in ay and "一案" in ay:
            m["案由"] = re.sub(r"^.*?诉.*?的?", "", ay).replace("一案", "").strip() or m.get("案由", "")
    case_no = sanitize_court_case_no(
        m.get("法院收案号") or m.get("案号") or "",
        pdf_text,
        case_type=case_type,
    )
    if case_no:
        m["法院收案号"] = case_no
        m["案号"] = case_no
    docs_raw = m.get("法院文件清单") or m.get("法院文书") or ""
    doc_list = parse_court_document_list(str(docs_raw))
    if doc_list:
        m["法院文件清单"] = "、".join(doc_list)
    # 民事/行政执行案件才使用执行终本专用合成逻辑；刑事案件保留裁判结果原意。
    if case_type in ("civil", "admin"):
        if doc_spans is not None and page_texts_by_path:
            m = build_outcome_from_units(m, doc_spans, page_texts_by_path, pdf_text)
        m = ensure_outcome_covers_execution(m, pdf_text, units=doc_spans)
    m = unify_case_outcome_fields(m)
    warnings = detect_outcome_warnings(m, doc_spans, pdf_text) if case_type in ("civil", "admin") else []
    if warnings:
        m["_outcome_warnings"] = warnings
    return sanitize_all_field_values(m)


def fill_all_templates(field_data, output_dir, log=print, output_options=None):
    os.makedirs(output_dir, exist_ok=True)
    generated = []
    layout_issues = []
    from output_options import templates_to_fill

    names = templates_to_fill(output_options)
    paths = get_template_paths()
    for name in names:
        template_path = paths.get(name)
        if not template_path or not os.path.exists(template_path):
            log(f"  [SKIP] 模板不存在: {name} → {template_path}")
            continue
        out_path = os.path.join(output_dir, f"{name}.docx")
        mapped = expand_fields_for_template(name, field_data)
        log(f"  [FILL] {name} ({len(mapped)} 字段)")
        filler = TemplateFiller(template_path)
        result = filler.fill_template(mapped, out_path, template_name=name)
        if isinstance(result, tuple):
            _, issues = result
            layout_issues.extend(issues or [])
        generated.append({"template": name, "path": out_path, "filename": f"{name}.docx"})
    return generated, layout_issues


def verify_outputs(output_dir):
    from docx import Document

    issues = []
    critical_fields = ["委托人", "当事人", "对方当事人", "案由", "审理法院", "承办律师", "收案日期"]

    for fname in os.listdir(output_dir):
        if not fname.endswith(".docx"):
            continue
        path = os.path.join(output_dir, fname)
        text = ""
        doc = Document(path)
        for t in doc.tables:
            for row in t.rows:
                for c in row.cells:
                    text += c.text
        for p in doc.paragraphs:
            text += p.text

        # 检查残留占位符
        left = re.findall(r"【([^】]{1,40})】", text)
        if left:
            critical_left = [f for f in left if f in critical_fields]
            if critical_left:
                issues.append(f"{fname}: 【严重】残留关键字段占位符 {len(critical_left)} 处: {', '.join(critical_left)}")
            else:
                issues.append(f"{fname}: 残留 {len(left)} 处占位符: {', '.join(set(left))}")

        # 检查关键字段是否为空或"待确认"
        for field in critical_fields:
            # 检查表格中是否有空的关键字段
            field_empty_pattern = f"{field}[:：]\\s*[\\n\\r]*$|{field}[:：]\\s*[\\n\\r]*待确认"
            if re.search(field_empty_pattern, text):
                issues.append(f"{fname}: 【警告】关键字段为空或待确认: {field}")

    return issues


_word_lock = threading.Lock()
WD_EXPORT_FORMAT_PDF = 17


def _word_convert_pairs(pairs, log=print):
    """逐文件独立 Word 进程转换 docx→pdf（避免批量 RPC 失败）"""
    import pythoncom
    import win32com.client

    results = {docx: False for docx, _ in pairs}
    if not pairs:
        return results

    def _convert_one(docx_path, pdf_path):
        pythoncom.CoInitialize()
        word = None
        ok = False
        err = None
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0
            doc = word.Documents.Open(os.path.abspath(docx_path))
            try:
                doc.ExportAsFixedFormat(
                    OutputFileName=os.path.abspath(pdf_path),
                    ExportFormat=WD_EXPORT_FORMAT_PDF,
                )
            except Exception as ex_export:
                doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=WD_EXPORT_FORMAT_PDF)
                err = f"export_fallback:{ex_export}"
            doc.Close(False)
            ok = os.path.exists(pdf_path)
        except Exception as e:
            err = str(e)
        finally:
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()
        return ok

    with _word_lock:
        for docx_path, pdf_path in pairs:
            log(f"  Word 转换: {os.path.basename(docx_path)}")
            results[docx_path] = _convert_one(docx_path, pdf_path)

    return results


def docx_to_pdf(docx_path, pdf_path, log=print):
    return _word_convert_pairs([(docx_path, pdf_path)], log=log).get(docx_path, False)


def merge_archive_pdf(docx_files, output_pdf):
    import fitz

    merger = fitz.open()
    for docx in docx_files:
        tmp_pdf = docx.replace(".docx", "_tmp.pdf")
        if docx_to_pdf(docx, tmp_pdf):
            src = fitz.open(tmp_pdf)
            merger.insert_pdf(src)
            src.close()
            try:
                os.remove(tmp_pdf)
            except OSError:
                pass
    if merger.page_count > 0:
        merger.save(output_pdf)
        merger.close()
        return True
    return False


def create_zip(source_dir, zip_path):
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for root, _, files in os.walk(source_dir):
            for f in files:
                if f.endswith((".docx", ".pdf", ".json")):
                    full = os.path.join(root, f)
                    arc = os.path.relpath(full, source_dir)
                    zf.write(full, arc)


# ============ V4 完整归档功能 ============

# 仅凭 catalog_seq（类型未对上）判定 found 所需的最低置信度；
# 低于此值视为「待复核」，不静默计入 found（会进入缺失清单提示人工核对）。
_FOUND_SEQ_CONF_THRESHOLD = 0.75


def _catalog_item_matches_span(item, span) -> bool:
    """判断一个 DocumentUnit 是否能满足目录项的 PDF/manual 子项。

    强匹配（doc_type 命中 doc_types 或 manual 映射）直接 found；
    仅 catalog_seq 命中而类型未对上时，要求切分置信度 >= 阈值，
    避免低置信误判为「已找到」。
    """
    doc_type = getattr(span, "doc_type", None)
    # 强匹配：类型命中
    if doc_type and doc_type in (item.doc_types or ()):
        return True
    manual_doc_type = ac.MANUAL_KEY_DOC_TYPES.get(item.manual_key)
    if manual_doc_type and doc_type == manual_doc_type:
        return True
    # 弱匹配：仅 catalog_seq 一致（类型未对上）→ 需置信度达标
    if getattr(span, "catalog_seq", None) == item.seq:
        conf = getattr(span, "confidence", 1.0)
        if conf is None:
            conf = 1.0
        return conf >= _FOUND_SEQ_CONF_THRESHOLD
    return False


def compute_found_seqs(catalog, doc_spans, generated_templates):
    """计算已找到的目录序号（纯函数，便于测试）

    Args:
        catalog: 目录列表
        doc_spans: DocSpan 列表
        generated_templates: 已生成的模板字典 {模板名: 路径}

    Returns:
        Set[int]: 已找到的目录序号集合
    """
    found_seqs = set()

    for item in catalog:
        is_found = False

        if item.source == "system":
            if any(t in generated_templates for t in item.templates):
                is_found = True
        elif any(_catalog_item_matches_span(item, span) for span in doc_spans):
            is_found = True

        if is_found:
            found_seqs.add(item.seq)

    return found_seqs


def compute_missing_items(case_type: str, catalog, found_seqs) -> List[Dict]:
    """按有效目录（可选 seq15 等）计算缺失项。"""
    import archive_catalog as ac

    found_set = set(found_seqs)
    effective = ac.get_effective_catalog(case_type, found_set)
    missing_items = []
    for item in effective:
        if item.seq not in found_set:
            missing_items.append({
                "seq": item.seq,
                "name": item.name,
                "source": item.source,
                "doc_types": item.doc_types,
                "manual_key": item.manual_key,
            })
    return missing_items


@dataclass
class ArchiveAnalysis:
    """归档分析结果（阶段1：analyze_archive）"""
    case_type: str  # 案件类型代码
    original_pdf: Optional[str]  # 原始 PDF 路径
    fields: Dict  # 提取的字段数据
    generated_templates: Dict[str, str]  # {模板名: docx路径}
    doc_spans: List  # DocSpan 列表
    found_seqs: Set[int]  # 已找到的目录序号
    missing_items: List[Dict]  # 缺失项列表 [{"seq": int, "name": str, "source": str}, ...]
    low_confidence_items: List[Dict] = field(default_factory=list)  # J-2 低置信度切分段
    template_issues: List[str] = field(default_factory=list)  # WF4 系统模板残留占位符/空字段
    outcome_warnings: List[str] = field(default_factory=list)  # V6 审办结果低置信预警


def _normalize_archive_sources(
    original_pdf: Optional[str] = None,
    sources: Optional[List] = None,
) -> List[DocumentSource]:
    """统一为 DocumentSource 列表（路径 A 单卷 / 路径 B 多文件）"""
    if sources:
        out = []
        for s in sources:
            if isinstance(s, DocumentSource):
                out.append(s)
            elif isinstance(s, dict):
                out.append(
                    DocumentSource(
                        path=s.get("path", ""),
                        doc_type=s.get("doc_type") or DOC_TYPE_DEFAULT,
                    )
                )
        return out
    if original_pdf:
        return [DocumentSource(path=original_pdf, doc_type=DOC_TYPE_DEFAULT)]
    return []


def ingest_archive_sources(
    doc_sources: List[DocumentSource],
    config: dict,
    log=print,
) -> tuple:
    """WF1：统一 OCR 摄入，每个源 PDF 至多一次重型 OCR。

    Returns:
        (pdf_texts_by_path, page_texts_by_path, layout_blocks_by_path,
         ocr_engine_calls, rapidocr_fallback_pages)
        ocr_engine_calls 仅计重型 OCR（MinerU/Baidu），不含 RapidOCR 单页回退。
    """
    from ocr_pipeline import ingest_pdf

    pdf_texts: Dict[str, str] = {}
    page_texts_map: Dict[str, List[str]] = {}
    layout_map: Dict[str, List[dict]] = {}
    total_calls = 0
    total_rapid = 0
    seen = set()

    for src in doc_sources:
        path = src.path
        if not path or path in seen or not os.path.exists(path):
            continue
        seen.add(path)
        log(f"       WF1: {os.path.basename(path)}")
        result = ingest_pdf(path, config, log=log)
        total_calls += result.ocr_engine_calls
        total_rapid += result.rapidocr_fallback_pages
        if result.rapidocr_fallback_pages:
            log(
                f"       RapidOCR 单页回退: {result.rapidocr_fallback_pages} 页"
                f"（不计入重型 OCR）"
            )
        if result.full_text:
            pdf_texts[path] = result.full_text
        if result.page_texts:
            page_texts_map[path] = result.page_texts
        if result.layout_blocks:
            layout_map[path] = result.layout_blocks

    return pdf_texts, page_texts_map, layout_map, total_calls, total_rapid


def segment_and_map_documents(
    doc_sources: List[DocumentSource],
    case_type: str,
    config: dict,
    *,
    pdf_texts_by_path: Optional[Dict[str, str]] = None,
    page_texts_by_path: Optional[Dict[str, List[str]]] = None,
    layout_blocks_by_path: Optional[Dict[str, List[dict]]] = None,
    log=print,
) -> List:
    """WF2+WF3：文书切分与目录映射"""
    return pdl.build_units_from_sources(
        doc_sources,
        case_type,
        config,
        pdf_texts=pdf_texts_by_path,
        page_texts_by_path=page_texts_by_path,
        layout_blocks_by_path=layout_blocks_by_path,
        log=log,
    )


def generate_system_templates(
    catalog,
    fields: Dict,
    log=print,
    *,
    issues_out: Optional[List[str]] = None,
    work_dir: Optional[str] = None,
) -> Dict[str, str]:
    """WF4：根据字段生成 source=system 的 docx 模板（失败不抛异常）

    生成后调用 verify_outputs 检测残留占位符/空关键字段；
    问题写入 issues_out（若提供）供上层（GUI/报告）提示，并打印日志。
    """
    generated_templates: Dict[str, str] = {}
    if not fields:
        log("       [WARN] 字段为空，仍生成系统模板（占位符可能保留，需人工补填）")
        log("       [HINT] 若刚部署 V5，请确认「系统设置」已填写 DeepSeek API Key")

    try:
        output_root = get_outputs_dir()
        if work_dir:
            os.makedirs(work_dir, exist_ok=True)
        else:
            work_dir = os.path.join(
                output_root, f"_analyze_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            )
            os.makedirs(work_dir, exist_ok=True)

        for item in catalog:
            if item.source != "system" or not item.templates:
                continue
            template_name = item.templates[0]
            try:
                paths = get_template_paths()
                template_path = paths.get(template_name)
                if not template_path or not os.path.exists(template_path):
                    continue
                out_path = os.path.join(work_dir, f"{template_name}.docx")
                mapped = expand_fields_for_template(template_name, fields)
                filler = TemplateFiller(template_path)
                filler.fill_template(mapped, out_path, template_name=template_name)
                generated_templates[template_name] = out_path
                log(f"       已生成: {template_name}")
            except Exception as e:
                log(f"       [WARN] 生成 {template_name} 失败: {e}")

        # 生成后校验：残留占位符 / 空关键字段
        if generated_templates:
            try:
                tpl_issues = verify_outputs(work_dir)
                if tpl_issues:
                    log(f"       [WARN] 系统模板校验发现 {len(tpl_issues)} 项问题:")
                    for it in tpl_issues:
                        log(f"          - {it}")
                    if issues_out is not None:
                        issues_out.extend(tpl_issues)
            except Exception as e:
                log(f"       [WARN] 系统模板校验失败: {e}")
    except Exception as e:
        log(f"       [WARN] 模板生成异常: {e}")

    return generated_templates


def extract_fields_from_text(pdf_text: str, log=print, case_type=None) -> Dict:
    """WF4 辅助：全文 → 归一化字段"""
    if not pdf_text:
        return {}
    try:
        raw_fields = extract_fields_auto(pdf_text, log=log, case_type=case_type)
        return normalize_fields(raw_fields, pdf_text, case_type=case_type)
    except Exception as e:
        log(f"       [WARN] 字段提取失败: {e}")
        return {}


def analyze_archive(
    case_type: str,
    original_pdf: Optional[str] = None,
    config: Optional[Dict] = None,
    *,
    sources: Optional[List] = None,
    log=print,
) -> ArchiveAnalysis:
    """分析案件归档（阶段1）

    执行（WF1~WF4）：
    1. WF1 统一 OCR 摄入（消除双 OCR）
    2. WF2+3 文书切分与目录映射
    3. WF4 字段提取 + 系统模板（可失败，不阻断切分）
    4. 计算缺失项（missing_items）

    Args:
        case_type: 案件类型代码（civil/criminal/admin/nonlit/counsel）
        original_pdf: 原始 PDF 路径（路径 A，与 sources 二选一）
        sources: DocumentSource 列表（路径 A/B）
        config: 配置字典（若为 None 则自动加载）
        log: 日志函数

    Returns:
        ArchiveAnalysis 对象
    """
    if config is None:
        config = load_config()

    from settings import pop_config, push_config

    push_config(config)
    try:
        return _analyze_archive_impl(
            case_type,
            original_pdf,
            config,
            sources=sources,
            log=log,
        )
    finally:
        pop_config()


def _analyze_archive_impl(
    case_type: str,
    original_pdf: Optional[str] = None,
    config: Optional[Dict] = None,
    *,
    sources: Optional[List] = None,
    log=print,
) -> ArchiveAnalysis:
    """analyze_archive 实际实现（在 push_config 作用域内调用）。"""
    if case_type not in ac.CASE_TYPE_LABELS:
        raise ValueError(f"Unknown case_type: {case_type}. Valid: {list(ac.CASE_TYPE_LABELS.keys())}")

    doc_sources = _normalize_archive_sources(original_pdf, sources)
    if not doc_sources:
        raise ValueError("analyze_archive 需要 original_pdf 或 sources")

    primary_pdf = next(
        (s.path for s in doc_sources if s.doc_type == DOC_TYPE_DEFAULT),
        doc_sources[0].path,
    )

    # 获取目录
    catalog = ac.get_catalog(case_type)
    log(f"分析案件类型: {ac.CASE_TYPE_LABELS[case_type]}（{len(catalog)} 项目录）")
    log(f"       输入文件: {len(doc_sources)} 个")

    log("[1/4] WF1 统一 OCR 摄入...")
    pdf_texts_by_path, page_texts_by_path, layout_blocks_by_path, ocr_calls, rapid_pages = (
        ingest_archive_sources(doc_sources, config, log=log)
    )
    log(f"       OCR 引擎调用: {ocr_calls} 次")
    if rapid_pages:
        log(f"       RapidOCR 单页回退: {rapid_pages} 页")

    log("[2/4] WF2+3 文书切分与映射...")
    doc_spans = []
    try:
        doc_spans = pdl.build_units_from_sources(
            doc_sources,
            case_type,
            config,
            pdf_texts=pdf_texts_by_path,
            page_texts_by_path=page_texts_by_path,
            layout_blocks_by_path=layout_blocks_by_path,
            log=log,
        )
        for u in doc_spans:
            n = u.end_page - u.start_page + 1
            seq_info = f" → seq{u.catalog_seq}" if u.catalog_seq is not None else ""
            try:
                log(
                    f"       {DOC_TYPE_LABELS.get(u.doc_type, u.doc_type)}: "
                    f"{os.path.basename(u.source_path)} 页{u.start_page}-{u.end_page} ({n}页){seq_info}"
                )
            except (UnicodeEncodeError, OSError):
                pass
    except Exception as e:
        import traceback
        log(f"       [WARN] 文书定位失败: {e}")
        log(traceback.format_exc())

    log("[3/4] WF4 字段提取与系统表...")
    pdf_text = "\n\n".join(pdf_texts_by_path.values())
    # 路径 B：多文件已带 doc_type，按 source 分路保留分类信息；
    # 路径 A：单卷（全部 default）仍按锚点切分（extract_fields_from_text）。
    has_typed_sources = any(s.doc_type != DOC_TYPE_DEFAULT for s in doc_sources)
    fields = {}
    if has_typed_sources:
        source_texts: Dict[str, str] = {}
        for src in doc_sources:
            t = pdf_texts_by_path.get(src.path, "")
            if t:
                source_texts[src.doc_type] = (
                    source_texts.get(src.doc_type, "") + "\n\n" + t
                ).strip()
        try:
            segmented = build_segmented_text(source_texts=source_texts).segments
            raw_fields = extract_fields_auto(pdf_text, segmented=segmented, log=log, case_type=case_type)
            fields = normalize_fields(
                raw_fields, pdf_text, doc_spans=doc_spans, page_texts_by_path=page_texts_by_path, case_type=case_type
            )
        except Exception as e:
            log(f"       [WARN] 分路字段提取失败: {e}")
            fields = {}
    else:
        segmented = build_segmented_from_units(doc_spans, page_texts_by_path)
        if segmented and any(segmented.values()):
            try:
                raw_fields = extract_fields_auto(
                    pdf_text, segmented=segmented, log=log, case_type=case_type
                )
                fields = normalize_fields(
                    raw_fields, pdf_text, doc_spans=doc_spans, page_texts_by_path=page_texts_by_path, case_type=case_type
                )
            except Exception as e:
                log(f"       [WARN] 分路字段提取失败: {e}")
                fields = extract_fields_from_text(pdf_text, log=log, case_type=case_type)
        else:
            fields = extract_fields_from_text(pdf_text, log=log, case_type=case_type)
    outcome_warnings = list(fields.pop("_outcome_warnings", None) or [])
    if outcome_warnings:
        for w in outcome_warnings:
            log(f"       [WARN] 审办结果: {w}")
    if fields:
        log(f"       提取到 {len(fields)} 个字段")
    elif not pdf_text:
        log("       [SKIP] 无 PDF 文本，跳过字段提取")
    else:
        log("       [WARN] 字段提取结果为空")
        ds_key = (config or {}).get("deepseek", {}).get("api_key", "")
        if not str(ds_key).strip():
            log("       [HINT] DeepSeek API Key 未配置，请到 V5「系统设置」填写")
        else:
            log("       [HINT] 请检查 DeepSeek 网络连通性与 OCR 文本质量")

    template_issues: List[str] = []
    preview_only = bool((config or {}).get("output", {}).get("preview_only"))
    if preview_only:
        generated_templates = {}
        log("       [PREVIEW] 当前为跨平台预览模式，跳过 DOCX 模板生成")
    else:
        generated_templates = generate_system_templates(
            catalog, fields, log=log, issues_out=template_issues
        )

    log("[4/4] 缺失项核对...")
    # 计算缺失项（V6：无执行案不将 seq15 计为缺失）
    catalog_templates = generated_templates
    if preview_only:
        # Browser previews replace the five system DOCX files in this mode.
        # Use virtual template names only for catalog completeness; they are
        # deliberately not persisted as file paths or passed to PDF assembly.
        catalog_templates = {
            item.templates[0]: "preview-only"
            for item in catalog
            if item.source == "system" and item.templates
        }
    found_seqs = compute_found_seqs(catalog, doc_spans, catalog_templates)
    missing_items = compute_missing_items(case_type, catalog, found_seqs)

    log(f"       已找到: {len(found_seqs)} 项")
    log(f"       缺失: {len(missing_items)} 项")

    try:
        from page_classifier import collect_low_confidence_units
        low_confidence_items = collect_low_confidence_units(doc_spans)
    except ImportError:
        low_confidence_items = []
    if low_confidence_items:
        log(f"       低置信切分: {len(low_confidence_items)} 段（可在 GUI 核对）")

    return ArchiveAnalysis(
        case_type=case_type,
        original_pdf=primary_pdf,
        fields=fields,
        generated_templates=generated_templates,
        doc_spans=doc_spans,
        found_seqs=found_seqs,
        missing_items=missing_items,
        low_confidence_items=low_confidence_items,
        template_issues=template_issues,
        outcome_warnings=outcome_warnings,
    )


def recompute_found_and_missing(analysis: "ArchiveAnalysis") -> "ArchiveAnalysis":
    """根据当前 doc_spans/generated_templates 重算 found_seqs 与 missing_items。

    用于用户在 GUI 手动调整文书归属（catalog_seq）后刷新缺失清单。
    原地更新并返回同一 analysis 对象。
    """
    import archive_catalog as ac

    catalog = ac.get_catalog(analysis.case_type)
    found_seqs = compute_found_seqs(catalog, analysis.doc_spans, analysis.generated_templates)
    analysis.found_seqs = found_seqs
    analysis.missing_items = compute_missing_items(analysis.case_type, catalog, found_seqs)
    return analysis


def apply_adjustments(analysis, adjustments, log=print):
    """应用手动调整（CLI 与 GUI 调序等价）后重算缺失清单。

    adjustments：列表，元素形如
        {"doc_id": 0, "catalog_seq": 14, "order": 1}
    - catalog_seq：可选，改文书目录归属；
    - order：可选，重排插入顺序（按 order 升序重写 doc_id）。
    未在 adjustments 中出现的文书保持原相对顺序，排在调整项之后。
    """
    if isinstance(adjustments, dict):
        adjustments = adjustments.get("adjustments", [])
    by_id = {}
    for a in adjustments or []:
        try:
            by_id[int(a["doc_id"])] = a
        except (KeyError, TypeError, ValueError):
            continue

    units = list(getattr(analysis, "doc_spans", None) or [])
    for u in units:
        a = by_id.get(getattr(u, "doc_id", None))
        if a and a.get("catalog_seq") is not None:
            try:
                u.catalog_seq = int(a["catalog_seq"])
            except (TypeError, ValueError):
                pass

    # 排序：有 order 的优先按 order，其余保持原顺序排后
    def sort_key(item):
        idx, u = item
        a = by_id.get(getattr(u, "doc_id", None))
        if a and a.get("order") is not None:
            try:
                return (0, int(a["order"]), idx)
            except (TypeError, ValueError):
                return (1, 0, idx)
        return (1, 0, idx)

    ordered = [u for _, u in sorted(enumerate(units), key=sort_key)]
    for new_id, u in enumerate(ordered):
        u.doc_id = new_id
    analysis.doc_spans = ordered

    recompute_found_and_missing(analysis)
    log(f"       已应用调整：{len(by_id)} 项，缺失 {len(analysis.missing_items)} 项")
    return analysis


def write_archive_report(analysis, result, output_pdf: str, log=print) -> Optional[str]:
    """将完整归档结果写为结构化 JSON 报告，便于人工核对与自动化。

    内容：成功标志、页数、源页守恒、缺失/跳过清单、排序问题、
    文书切分摘要、低置信段、系统模板问题。返回报告路径（失败返回 None）。
    """
    try:
        base = os.path.splitext(output_pdf)[0]
        report_path = f"{base}_archive_report.json"

        doc_spans_summary = []
        for u in (getattr(analysis, "doc_spans", None) or []):
            doc_spans_summary.append({
                "doc_id": getattr(u, "doc_id", None),
                "catalog_seq": getattr(u, "catalog_seq", None),
                "doc_type": getattr(u, "doc_type", None),
                "source": os.path.basename(getattr(u, "source_path", "") or ""),
                "start_page": getattr(u, "start_page", None),
                "end_page": getattr(u, "end_page", None),
                "confidence": round(float(getattr(u, "confidence", 1.0) or 1.0), 3),
            })

        report = {
            "output_pdf": getattr(result, "output_pdf", output_pdf),
            "success": bool(getattr(result, "success", False)),
            "page_count": getattr(result, "page_count", 0),
            "original_pages_included": getattr(result, "original_pages_included", 0),
            "missing": getattr(result, "missing", None) or [],
            "order_issues": getattr(result, "order_issues", None) or [],
            "doc_spans": doc_spans_summary,
            "low_confidence_items": getattr(analysis, "low_confidence_items", None) or [],
            "template_issues": getattr(analysis, "template_issues", None) or [],
            "case_type": getattr(analysis, "case_type", None),
        }

        import json as _json
        with open(report_path, "w", encoding="utf-8") as f:
            _json.dump(report, f, ensure_ascii=False, indent=2)
        log(f"       归档报告: {report_path}")
        return report_path
    except Exception as e:
        log(f"       [WARN] 写归档报告失败: {e}")
        return None


# ArchiveResult 从 pdf_archive_merger 导入，避免重复定义
from pdf_archive_merger import ArchiveResult


def assemble_archive(
    analysis: ArchiveAnalysis,
    output_pdf: str,
    *,
    supplement_files: Optional[List[str]] = None,
    supplements: Optional[Dict[int, List[str]]] = None,
    skipped: Optional[List[int]] = None,
    config: Optional[Dict] = None,
    log=print,
) -> ArchiveResult:
    """拼装完整归档 PDF（阶段2）

    Args:
        analysis: analyze_archive 的返回结果
        output_pdf: 输出 PDF 路径
        supplement_files: 用户补充的原始文件列表（将调用 classify_attachments 分类）
        supplements: 已分类的附件 {seq: [文件路径]}（若提供则跳过分类）
        skipped: 用户跳过的 seq 列表
        config: 配置字典（用于附件分类）
        log: 日志函数

    Returns:
        ArchiveResult
    """
    try:
        import pdf_archive_merger as pam
        import attachment_classifier as attc
        import archive_catalog as ac
    except ImportError:
        log("必要模块未找到")
        return ArchiveResult(output_pdf, success=False, missing=[])

    if config is None:
        config = load_config()

    # 若提供了原始文件列表，调用 classify_attachments 分类
    if supplement_files and not supplements:
        log(f"分类 {len(supplement_files)} 个补充附件...")
        classified = attc.classify_attachments(
            supplement_files, analysis.case_type, config, log=log
        )
        # 按 seq 分组
        supplements = {}
        for att in classified:
            if att.catalog_item is not None:
                seq = att.catalog_item.seq
            else:
                # 未匹配到目录项的归入 evidence 槽（使用 manual_key 查找）
                evidence_item = ac.catalog_item_for_manual_key(analysis.case_type, "evidence")
                seq = evidence_item.seq if evidence_item else -1
            if seq >= 0 and seq not in supplements:
                supplements[seq] = []
            if seq >= 0:
                supplements[seq].append(att.file_path)

    if supplements is None:
        supplements = {}
    if skipped is None:
        skipped = []

    order_mode = get_archive_order_mode(config)
    log(f"开始拼装完整归档 PDF...（排序模式: {order_mode}）")

    # 调用 pdf_archive_merger 的核心函数
    result = pam.build_full_archive(
        case_type=analysis.case_type,
        original_pdf=analysis.original_pdf,
        generated_templates=analysis.generated_templates,
        doc_spans=analysis.doc_spans,
        supplements=supplements,
        skipped=skipped,
        output_pdf=output_pdf,
        log=log,
        order_mode=order_mode,
    )

    if result.success:
        log(f"归档完成: {result.page_count} 页")
        log(f"缺失/跳过项: {len(result.missing)} 项")
    else:
        log("归档失败")

    return ArchiveResult(
        output_pdf=result.output_pdf,
        success=result.success,
        missing=result.missing,
        page_count=result.page_count,
        sources=result.sources,
        original_pages_included=getattr(result, "original_pages_included", 0),
        order_issues=getattr(result, "order_issues", None),
    )


def _resolve_max_pages(config, pdf_path, max_pages=None):
    total_pages = get_pdf_page_count(pdf_path)
    cfg_mp = config.get("local_ocr", {}).get("max_pages", 0)
    if max_pages is not None:
        use_pages = max_pages
    elif cfg_mp and cfg_mp > 0:
        use_pages = cfg_mp
    else:
        use_pages = total_pages or cfg_mp
    if total_pages > 0 and use_pages > total_pages:
        use_pages = total_pages
    if use_pages <= 0 and total_pages > 0:
        use_pages = total_pages
    config.setdefault("local_ocr", {})["max_pages"] = use_pages
    return use_pages, total_pages


def _derive_case_key(field_data: dict, fallback: str) -> str:
    for key in ("法院收案号", "案号", "案件或项目名称"):
        val = (field_data.get(key) or "").strip()
        if val:
            safe = re.sub(r'[<>:"/\\|?*]', "_", val)[:40]
            return safe
    return fallback


def _finalize_archive(
    work_dir,
    base_name,
    field_data,
    pdf_text,
    generated,
    layout_issues,
    docx_only,
    output_root,
    total_steps,
    log,
    output_options=None,
):
    from output_options import normalize_output_options

    opts = normalize_output_options(output_options)
    issues = verify_outputs(work_dir)

    critical_issues = [i for i in issues if "【严重】" in i or "【警告】" in i]
    normal_issues = [i for i in issues if "【严重】" not in i and "【警告】" not in i]

    if critical_issues:
        log("  [ERROR] 发现严重问题，建议检查 LLM 提取或字段映射:")
        for i in critical_issues:
            log(f"    ❌ {i}")

    for i in normal_issues:
        log(f"  [WARN] {i}")

    try:
        from layout_verify import verify_output_layout

        extra = verify_output_layout(work_dir, field_data, log=log)
        for e in extra:
            if e not in layout_issues:
                layout_issues.append(e)
    except Exception as ex:
        log(f"  [WARN] 版式校验跳过: {ex}")

    zip_path = None
    archive_pdf = None

    if docx_only:
        log(f"[4/{total_steps}] 完成：已生成 {len(generated)} 份 docx")
        for g in generated:
            log(f"       · {g['filename']}")
    else:
        log(f"[4/{total_steps}] 校验输出...")
        log(f"[5/{total_steps}] 打包归档...")
        zip_path = os.path.join(work_dir, f"{base_name}_归档资料.zip")
        create_zip(work_dir, zip_path)
        docx_list = [g["path"] for g in generated]
        archive_pdf = os.path.join(work_dir, f"{base_name}_归档资料.pdf")
        if merge_archive_pdf(docx_list, archive_pdf):
            log(f"       合并 PDF: {archive_pdf}")
        else:
            archive_pdf = None
            log("  [WARN] 合并 PDF 失败（需本机安装 Word），仍可使用 ZIP 内 docx")
        for src in (zip_path, archive_pdf):
            if src and os.path.exists(src):
                dest = os.path.join(output_root, os.path.basename(src))
                shutil.copy2(src, dest)

    return {
        "success": True,
        "output_dir": work_dir,
        "zip_path": zip_path,
        "archive_pdf": archive_pdf,
        "generated_files": generated,
        "output_mode": opts["mode"],
        "field_count": len(field_data),
        "fields": field_data,
        "text_length": len(pdf_text),
        "verify_issues": issues,
        "layout_issues": layout_issues,
        "docx_only": docx_only,
    }


def _ocr_engine_label(engine: str) -> str:
    return {
        "mineru": "MinerU 本地",
        "mineru_api": "MinerU API",
        "baidu": "百度 OCR",
    }.get(engine, engine)


def process_archive_sources(
    sources: list,
    output_dir=None,
    prompt_path=None,
    max_pages=None,
    log=print,
    output_options=None,
):
    """多 PDF 分类归档（单案件）"""
    doc_sources = []
    for s in sources:
        if isinstance(s, DocumentSource):
            doc_sources.append(s)
        elif isinstance(s, dict):
            doc_sources.append(
                DocumentSource(
                    path=s.get("path", ""),
                    doc_type=s.get("doc_type") or DOC_TYPE_DEFAULT,
                )
            )

    err = validate_sources_for_archive(doc_sources)
    if err:
        return {"success": False, "error": err}

    config = load_config()
    docx_only = output_docx_only(config)
    output_root = get_outputs_dir()
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    first_stem = os.path.splitext(os.path.basename(doc_sources[0].path))[0]
    work_dir = output_dir or os.path.join(output_root, f"{first_stem}_multi_{stamp}")
    os.makedirs(work_dir, exist_ok=True)

    engine = get_ocr_engine(config)
    eng_label = _ocr_engine_label(engine)
    total_steps = 4 if docx_only else 5

    source_texts = {}
    pdf_text_parts = []
    for i, src in enumerate(doc_sources, 1):
        if not os.path.isfile(src.path):
            return {"success": False, "error": f"PDF 不存在: {src.path}"}
        use_pages, total_pages = _resolve_max_pages(config, src.path, max_pages)
        log(
            f"[1/{total_steps}] OCR ({eng_label}) [{i}/{len(doc_sources)}] "
            f"{os.path.basename(src.path)} ({src.doc_type})"
            + (f" 共{total_pages}页" if total_pages else "")
        )
        text, ocr_err = extract_pdf_text(src.path, config, log=log)
        if not text or len(text.strip()) < 20:
            return {"success": False, "error": f"OCR 失败: {src.path} — {ocr_err or '文本过短'}"}
        source_texts[src.doc_type] = source_texts.get(src.doc_type, "") + "\n\n" + text
        pdf_text_parts.append(text)
        log(f"       已提取 {len(text)} 字符")

    pdf_text = "\n\n".join(pdf_text_parts)
    segmented = build_segmented_text(source_texts=source_texts).segments

    log(f"[2/{total_steps}] DeepSeek 分路提取字段...")
    try:
        raw_fields = extract_fields_auto(pdf_text, segmented=segmented, log=log)
    except Exception as e:
        return {"success": False, "error": f"字段提取失败: {e}"}
    field_data = normalize_fields(raw_fields, pdf_text)
    if not field_data:
        return {"success": False, "error": "未提取到任何字段"}
    log(f"       共 {len(field_data)} 个字段")

    if not docx_only:
        meta_path = os.path.join(work_dir, "extracted_fields.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(field_data, f, ensure_ascii=False, indent=2)

    from output_options import templates_to_fill, mode_label, normalize_output_options

    opts = normalize_output_options(output_options)
    fill_names = templates_to_fill(opts)
    log(f"[3/{total_steps}] 填充模板（{mode_label(opts['mode'])}，{len(fill_names)} 份）...")
    generated, layout_issues = fill_all_templates(
        field_data, work_dir, log=log, output_options=opts
    )
    if not generated:
        return {"success": False, "error": "未能生成任何文书"}

    base_name = _derive_case_key(field_data, first_stem)
    return _finalize_archive(
        work_dir,
        base_name,
        field_data,
        pdf_text,
        generated,
        layout_issues,
        docx_only,
        output_root,
        total_steps,
        log,
        output_options=opts,
    )


def process_archive(
    pdf_path,
    output_dir=None,
    prompt_path=None,
    max_pages=None,
    log=print,
    output_options=None,
):
    """一键归档：PDF → 5 份 docx（output.docx_only=true 时无 zip/pdf/json）"""
    if not os.path.exists(pdf_path):
        return {"success": False, "error": f"PDF 不存在: {pdf_path}"}

    config = load_config()
    docx_only = output_docx_only(config)
    output_root = get_outputs_dir()
    base_name = os.path.splitext(os.path.basename(pdf_path))[0]
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    work_dir = output_dir or os.path.join(output_root, f"{base_name}_{stamp}")
    os.makedirs(work_dir, exist_ok=True)
    use_pages, total_pages = _resolve_max_pages(config, pdf_path, max_pages)

    engine = get_ocr_engine(config)
    eng_label = _ocr_engine_label(engine)
    total_steps = 4 if docx_only else 5
    log(
        f"[1/{total_steps}] 提取 PDF 文本 ({eng_label}): {pdf_path}"
        + (f"（共 {total_pages} 页，解析 {use_pages} 页）" if total_pages else "")
    )
    pdf_text, ocr_err = extract_pdf_text(pdf_path, config, log=log)
    if not pdf_text or len(pdf_text.strip()) < 100:
        return {"success": False, "error": f"PDF 文本提取失败: {ocr_err or '文本过短'}"}
    log(f"       已提取 {len(pdf_text)} 字符")

    segmented_obj = build_segmented_text(pdf_text=pdf_text)
    log(f"[2/{total_steps}] DeepSeek 提取字段（{get_extraction_mode()}）...")
    try:
        raw_fields = extract_fields_auto(
            pdf_text, segmented=segmented_obj.segments, log=log
        )
    except Exception as e:
        return {"success": False, "error": f"字段提取失败: {e}"}
    field_data = normalize_fields(raw_fields, pdf_text)
    if not field_data:
        return {"success": False, "error": "未提取到任何字段"}
    log(f"       共 {len(field_data)} 个字段")

    if not docx_only:
        meta_path = os.path.join(work_dir, "extracted_fields.json")
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(field_data, f, ensure_ascii=False, indent=2)
        log(f"       字段缓存 → {meta_path}")

    from output_options import templates_to_fill, mode_label, normalize_output_options

    opts = normalize_output_options(output_options)
    fill_names = templates_to_fill(opts)
    log(f"[3/{total_steps}] 填充模板（{mode_label(opts['mode'])}，{len(fill_names)} 份）...")
    generated, layout_issues = fill_all_templates(
        field_data, work_dir, log=log, output_options=opts
    )
    if not generated:
        return {"success": False, "error": "未能生成任何文书（请检查 templates/bundled 下 .doc 模板）"}

    return _finalize_archive(
        work_dir,
        base_name,
        field_data,
        pdf_text,
        generated,
        layout_issues,
        docx_only,
        output_root,
        total_steps,
        log,
        output_options=opts,
    )
