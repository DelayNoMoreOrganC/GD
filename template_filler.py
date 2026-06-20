#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
法律文档模板填充 - 基于原始 .doc 模板，保留版式
使用 Word COM 在文档内替换【占位符】，避免破坏表格结构
"""

import os
import re
import shutil
import tempfile

from field_mapping import ORIGINAL_TEMPLATE_PATHS
from manifest_word_fill import fill_document_by_manifest, fill_seq_cells_by_manifest
from table_layout_optimizer import (
    optimize_document_tables,
    restore_table_row_heights,
    snapshot_table_row_heights,
)

# Word 常量
WD_REPLACE_ALL = 2
WD_FIND_CONTINUE = 1
WD_FORMAT_DOCX = 16
WD_COLOR_BLACK = 0
WD_COLOR_AUTO = -16777216
SEQ_PREFIX = "__seq__"


def _blacken_range(rng):
    """将范围内文字设为黑色（含新填入内容，覆盖模板红色），不加粗"""
    try:
        rng.Font.Color = WD_COLOR_BLACK
        rng.Font.ColorIndex = 1
        rng.Font.HighlightColorIndex = 0
        rng.Font.Bold = 0
    except Exception:
        pass
    try:
        from word_atomic_fill import clear_range_bold

        clear_range_bold(rng)
    except Exception:
        pass
    try:
        chars = rng.Characters
        n = chars.Count
        for i in range(1, n + 1):
            try:
                ch = chars(i)
                ch.Font.Color = WD_COLOR_BLACK
                try:
                    ch.Font.ColorIndex = 1
                except Exception:
                    pass
            except Exception:
                pass
    except Exception:
        try:
            runs = rng.Runs
            for i in range(1, runs.Count + 1):
                run = runs(i)
                run.Font.Color = WD_COLOR_BLACK
                try:
                    run.Font.ColorIndex = 1
                except Exception:
                    pass
        except Exception:
            pass


def _set_document_text_black(doc):
    """全文设为黑色"""
    try:
        _blacken_range(doc.Content)
    except Exception:
        pass
    try:
        for ti in range(1, doc.Tables.Count + 1):
            table = doc.Tables(ti)
            for ri in range(1, table.Rows.Count + 1):
                row = table.Rows(ri)
                for ci in range(1, row.Cells.Count + 1):
                    try:
                        _blacken_range(row.Cells(ci).Range)
                    except Exception:
                        pass
    except Exception:
        pass
    try:
        for story in range(1, doc.StoryRanges.Count + 1):
            try:
                _blacken_range(doc.StoryRanges(story))
            except Exception:
                pass
    except Exception:
        pass


def _force_document_black(doc):
    """Find 将红色字体批量改为黑色（模板说明字多为红色）"""
    wd_replace_all = 2
    wd_find_continue = 1
    color_indices = (6, 7, 13)
    font_colors = (255, 16711680, 0xFF0000)
    for color_index in color_indices:
        try:
            find = doc.Content.Find
            find.ClearFormatting()
            find.Font.ColorIndex = color_index
            find.Replacement.ClearFormatting()
            find.Replacement.Font.Color = WD_COLOR_BLACK
            find.Replacement.Font.ColorIndex = 1
            find.Replacement.Font.HighlightColorIndex = 0
            find.Replacement.Font.Bold = 0
            find.Execute(
                Replace=wd_replace_all,
                Forward=True,
                Wrap=wd_find_continue,
                Format=True,
            )
        except Exception:
            pass
    for font_color in font_colors:
        try:
            find = doc.Content.Find
            find.ClearFormatting()
            find.Font.Color = font_color
            find.Replacement.ClearFormatting()
            find.Replacement.Font.Color = WD_COLOR_BLACK
            find.Replacement.Font.ColorIndex = 1
            find.Replacement.Font.HighlightColorIndex = 0
            find.Replacement.Font.Bold = 0
            find.Execute(
                Replace=wd_replace_all,
                Forward=True,
                Wrap=wd_find_continue,
                Format=True,
            )
        except Exception:
            pass
    _set_document_text_black(doc)


def _split_field_data(field_data):
    normal = {}
    sequential = {}
    for key, val in (field_data or {}).items():
        if key.startswith(SEQ_PREFIX):
            name = key[len(SEQ_PREFIX) :]
            if isinstance(val, (list, tuple)):
                sequential[name] = [str(x).strip() for x in val if str(x).strip()]
            elif val:
                sequential[name] = [str(val).strip()]
        else:
            normal[key] = val
    return normal, sequential


def _apply_placeholders_to_text(text, fields):
    if not text:
        return text, 0
    hits = 0
    new_text = text
    for field_name, value in fields.items():
        if value is None:
            continue
        val = str(value)
        bracketed = f"【{field_name}】"
        if bracketed in new_text:
            cnt = new_text.count(bracketed)
            new_text = new_text.replace(bracketed, val)
            hits += cnt
        if field_name.startswith("=") and field_name in new_text:
            cnt = new_text.count(field_name)
            new_text = new_text.replace(field_name, val)
            hits += cnt
    return new_text, hits


def _word_fill_document(template_path, field_data, output_path, template_name=None):
    """用 Microsoft Word 打开模板并替换占位符（保留原始格式）"""
    import pythoncom
    import win32com.client

    try:
        from archive_pipeline import _word_lock
    except ImportError:
        from contextlib import nullcontext
        _word_lock = nullcontext()

    with _word_lock:
        pythoncom.CoInitialize()
        word = None
        doc = None
        replaced = 0
        try:
            word = win32com.client.DispatchEx("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(os.path.abspath(template_path))

            row_snapshot = snapshot_table_row_heights(doc)

            if template_name == "送达材料清单":
                from delivery_list_filler import word_fill_delivery_list

                word_fill_delivery_list(doc, field_data, output_path)
                doc.Close(False)
                layout_issues = []
                try:
                    from layout_verify import verify_template

                    layout_issues = verify_template("送达材料清单", output_path)
                    for issue in layout_issues:
                        print(f"  [LAYOUT] {issue}")
                except Exception:
                    pass
                return output_path, layout_issues

            if not template_name:
                raise ValueError("template_name 必填（V1.2 映射表填充）")

            normal_fields, sequential_fields = _split_field_data(field_data)

            replaced, coords = fill_document_by_manifest(
                doc,
                template_name,
                normal_fields,
                blacken_fn=_blacken_range,
                validate=False,
            )
            if sequential_fields:
                seq_n, _seq_coords = fill_seq_cells_by_manifest(
                    doc,
                    template_name,
                    sequential_fields,
                    blacken_fn=_blacken_range,
                )
                replaced += seq_n

            optimize_document_tables(doc, template_name=template_name)

            restored = restore_table_row_heights(doc, row_snapshot)
            if restored:
                print(f"  [OK] 已恢复 {restored} 行表格行高")

            from post_fill_cleanup import finalize_fill_document

            finalize_fill_document(
                doc,
                template_name,
                blacken_fn=_blacken_range,
                field_patch=normal_fields,
            )

            out_abs = os.path.abspath(output_path)
            os.makedirs(os.path.dirname(out_abs) or ".", exist_ok=True)
            if output_path.lower().endswith(".docx"):
                doc.SaveAs2(out_abs, FileFormat=WD_FORMAT_DOCX)
            else:
                doc.SaveAs2(out_abs)
            doc.Close(False)
            print(f"\n[OK] Word 填充完成，约 {replaced} 类占位符，输出: {output_path}")

            layout_issues = []
            try:
                from layout_verify import verify_template

                layout_issues = verify_template(template_name, output_path)
                for issue in layout_issues:
                    print(f"  [LAYOUT] {issue}")
            except Exception as e:
                print(f"  [WARN] 版式校验跳过: {e}")

            return output_path, layout_issues
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()


def _docx_replace_placeholders(docx_path, field_data, output_path):
    """无 Word 时的降级：在 docx 段落/单元格内替换【占位符】"""
    from docx import Document
    from docx.shared import RGBColor

    doc = Document(docx_path)
    filled = 0
    normal_fields, sequential_fields = _split_field_data(field_data)
    black = RGBColor(0, 0, 0)

    def replace_in_paragraph(paragraph, fields, one_ph=None, one_val=None):
        nonlocal filled
        if one_ph:
            ph = f"【{one_ph}】"
            if ph in paragraph.text:
                for run in paragraph.runs:
                    if ph in run.text:
                        run.text = run.text.replace(ph, one_val, 1)
                        filled += 1
            return
        if "【" not in paragraph.text and not any(k in paragraph.text for k in fields):
            return
        for field_name, value in fields.items():
            if value is None:
                continue
            ph = f"【{field_name}】"
            for run in paragraph.runs:
                if ph in run.text:
                    run.text = run.text.replace(ph, str(value))
                    filled += 1

    def blacken_paragraph(paragraph):
        for run in paragraph.runs:
            if run.text.strip():
                run.font.color.rgb = black
                run.font.bold = False

    for para in doc.paragraphs:
        replace_in_paragraph(para, normal_fields)
        blacken_paragraph(para)

    for table in doc.tables:
        for placeholder, values in sequential_fields.items():
            idx = 0
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        ph = f"【{placeholder}】"
                        if ph not in para.text:
                            continue
                        val = values[idx] if idx < len(values) else ""
                        replace_in_paragraph(para, {}, one_ph=placeholder, one_val=val)
                        blacken_paragraph(para)
                        idx += 1
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, normal_fields)
                    blacken_paragraph(para)

    doc.save(output_path)
    print(f"[OK] docx 占位符替换完成: {output_path}")
    return output_path


def convert_doc_to_docx(doc_path, docx_path=None):
    """Word 将 .doc 转为 .docx（正确转换，非改扩展名）"""
    import pythoncom
    import win32com.client

    if docx_path is None:
        docx_path = doc_path.replace(".doc", ".docx")

    if os.path.exists(docx_path) and os.path.getsize(docx_path) > 5000:
        return docx_path

    pythoncom.CoInitialize()
    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        doc = word.Documents.Open(os.path.abspath(doc_path))
        doc.SaveAs2(os.path.abspath(docx_path), FileFormat=WD_FORMAT_DOCX)
        doc.Close(False)
        print(f"[OK] 已转换: {docx_path}")
        return docx_path
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()


class TemplateFiller:
    """模板填充器 - 优先 Word COM，保留原始 .doc 版式"""

    def __init__(self, template_path):
        self.template_path = template_path
        self.work_path = template_path
        self._temp_docx = None

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板不存在: {template_path}")

        if template_path.lower().endswith(".doc"):
            cache_dir = os.path.join(
                os.path.dirname(__file__), "templates_local", "original"
            )
            os.makedirs(cache_dir, exist_ok=True)
            base = os.path.splitext(os.path.basename(template_path))[0]
            cache_name = None
            for name, path in ORIGINAL_TEMPLATE_PATHS.items():
                if os.path.abspath(path) == os.path.abspath(template_path):
                    cache_name = f"{name}.docx"
                    break
            if not cache_name:
                cache_name = re.sub(r'[<>:"/\\|?*]', "_", base) + ".docx"
            cached = os.path.join(cache_dir, cache_name)
            self.work_path = convert_doc_to_docx(template_path, cached)

    def fill_template(self, field_data, output_path, template_name=None):
        """填充模板并保存；返回 output_path 或 (path, layout_issues)"""
        source = self.template_path
        if source.lower().endswith(".doc") and os.path.exists(source):
            try:
                result = _word_fill_document(
                    source, field_data, output_path, template_name=template_name
                )
                if isinstance(result, tuple):
                    return result
                return result, []
            except Exception as e:
                print(f"  [WARN] Word 填充 .doc 失败: {e}，尝试 docx 降级")

        if self.work_path.lower().endswith(".docx"):
            try:
                result = _word_fill_document(
                    self.work_path, field_data, output_path, template_name=template_name
                )
                if isinstance(result, tuple):
                    return result
                return result, []
            except Exception as e:
                print(f"  [WARN] Word 填充 docx 失败: {e}，尝试 python-docx")

        return _docx_replace_placeholders(self.work_path, field_data, output_path), []


def demo_usage():
    from field_mapping import expand_fields_for_template

    sample = {
        "委托人": "佛山某银行股份有限公司",
        "当事人": "佛山某银行股份有限公司",
        "对方当事人": "李四",
        "委托人电话": "83039793",
        "地址": "广东省佛山市南海区",
        "收费标准": "固定5000元\\基础3000元+风险",
        "收案日期": "2019-09-01",
        "承办律师": "张三律师",
        "案由": "金融借款合同纠纷",
        "审理法院": "佛山市南海区人民法院",
        "审级": "一审、执行",
        "法院收案号": "（2019）粤0605民初xxx号",
        "案情简介": "李四的贷款逾期，佛山某银行委托我所代理起诉，起诉标的500000元",
        "结案小结": "判决支持原告诉请，案件已执行完毕",
        "法院文件清单": "民事判决书、执行裁定书",
    }

    for name, path in ORIGINAL_TEMPLATE_PATHS.items():
        if not os.path.exists(path):
            print(f"[SKIP] 不存在: {path}")
            continue
        print(f"\n{'=' * 60}\n处理: {name}")
        mapped = expand_fields_for_template(name, sample)
        filler = TemplateFiller(path)
        out = os.path.join("./outputs", f"{name}_demo_filled.docx")
        filler.fill_template(mapped, out)


if __name__ == "__main__":
    demo_usage()
